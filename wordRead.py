from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shape import InlineShape
import os
import olefile

print("bins")
data = ['test.png','test1.png']
i = 0
ole = olefile.OleFileIO("testing.docx")
# doc = docx.Document("testingDemo.docx")
doc = Document("testing.docx")
# imgType = "WD_INLINE_SHAPE.LINKED_PICTURE"
# for oleEmbed in doc.inline_shapes :
	# blip = doc.inline_shapes[i]._inline.graphic.graphicData.pic.blipFill.blip
	# rID = blip.embed
	# rID = blip.embed
	# document_part = doc.part
	# image_part = document_part.related_parts[rID]
	# fr = open(data[i], "wb")
	# fr.write(image_part._blob)
	# i += 1
	# doc.add_picture(R"AutomationProject\gitConfigList.png")
	# doc.save("testing.docx")
# fr.close()


print("ends")
